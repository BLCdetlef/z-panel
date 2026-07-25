#!/usr/bin/env python3
from __future__ import annotations

"""ZUSTAND News Studio 5.5.0 – Wissenschaftliche Redaktion.

Benötigt im selben Ordner:
- news_studio_5_3.py
- news_studio_5_2_1.py (sowie dessen bisherige Basisdateien)

Neue Funktionen:
- Eigene Funde (URL/LinkedIn/DOI/Text/PDF/Screenshot) lokal erfassen
- Kontaktstatus mit Historie
- Redaktionsstatistik
- sichere Entfernung von Artikel-Kontakt-Verknüpfungen
- Interviewanfragen mit Studiozeiten und transparentem Ablauf
- Dankeschön-Vorlage
- Interviewmappe als DOCX
- Qualitätsprüfung je Beitrag
"""

import importlib.util
import json
import os
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

SCRIPT_DIR = Path(__file__).resolve().parent
BASE_SCRIPT = SCRIPT_DIR / "news_studio_5_3.py"
if not BASE_SCRIPT.exists():
    raise SystemExit(
        "news_studio_5_3.py wurde nicht gefunden.\n"
        "Lege Version 5.4.2 in denselben Ordner wie Version 5.3."
    )

spec = importlib.util.spec_from_file_location("news_studio_5_3_base", BASE_SCRIPT)
if spec is None or spec.loader is None:
    raise SystemExit("News Studio 5.3 konnte nicht geladen werden.")
base53 = importlib.util.module_from_spec(spec)
sys.modules[spec.name] = base53
spec.loader.exec_module(base53)

tk = base53.tk
ttk = base53.ttk
messagebox = base53.messagebox
clean_text = base53.clean_text
write_contact_db = base53.write_contact_db

try:
    from tkinter import filedialog
except ImportError:  # pragma: no cover
    filedialog = None

STATUS_VALUES = (
    "Recherchiert",
    "E-Mail versendet",
    "Antwort erhalten",
    "Interview vereinbart",
    "Interview durchgeführt",
    "Beitrag veröffentlicht",
    "Dankeschön versendet",
)

QUALITY_FIELDS = (
    ("primarySource", "Primärquelle geprüft"),
    ("doiOrStableUrl", "DOI oder stabile Originalquelle vorhanden"),
    ("contactChecked", "Ansprechperson geprüft"),
    ("plainLanguage", "Kurztext verständlich und ohne Übertreibung"),
    ("imageFit", "Bild passt zum Inhalt"),
    ("sourcesComplete", "Quellen vollständig"),
    ("panelTested", "Darstellung im Infoscreen geprüft"),
    ("interviewPotential", "Interviewpotenzial bewertet"),
)


def now_iso() -> str:
    return datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")


def safe_slug(value: str, fallback: str = "fund") -> str:
    value = clean_text(value).lower()
    chars = []
    for char in value:
        if char.isalnum():
            chars.append(char)
        elif chars and chars[-1] != "-":
            chars.append("-")
    slug = "".join(chars).strip("-")
    return slug[:70] or fallback


def ensure_workflow_schema(db: dict[str, Any]) -> None:
    db.setdefault("workflowVersion", 1)
    db.setdefault("workflow", {})
    db.setdefault("qualityChecks", {})
    db.setdefault("ownFinds", {})
    for contact_id, contact in db.get("contacts", {}).items():
        if not isinstance(contact, dict):
            continue
        wf = db["workflow"].setdefault(contact_id, {})
        wf.setdefault("status", "Recherchiert")
        wf.setdefault("history", [])
        wf.setdefault("notes", "")


WORKFLOW_FILE = SCRIPT_DIR / "newsredaktion" / "kontakte" / "workflow.json"


def read_workflow_file() -> dict[str, Any]:
    """Liest die 5.4-Workflowdaten unabhängig von älteren Kontaktschemata."""
    if not WORKFLOW_FILE.exists():
        return {}
    try:
        data = json.loads(WORKFLOW_FILE.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError):
        return {}


def write_workflow_file(workflow: dict[str, Any]) -> None:
    """Speichert Kontaktstatus, Verlauf und Notizen atomar."""
    WORKFLOW_FILE.parent.mkdir(parents=True, exist_ok=True)
    temp = WORKFLOW_FILE.with_suffix(".json.tmp")
    temp.write_text(
        json.dumps(workflow, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    temp.replace(WORKFLOW_FILE)


EDITORIAL_FILE = SCRIPT_DIR / "newsredaktion" / "redaktion" / "redaktion.json"

DEFAULT_EDITORIAL_GUIDE = {
    "mission": (
        "ZUSTAND macht den Zustand unserer natürlichen Lebensgrundlagen sichtbar "
        "und verständlich. Grundlage sind wissenschaftliche Primärquellen, "
        "transparente Recherche, nachvollziehbare Zusammenhänge und der Dialog "
        "mit Forschenden. ZUSTAND möchte Orientierung geben, nicht polarisieren."
    ),
    "code": """1. Wir bevorzugen wissenschaftliche Primärquellen.\n2. Wir unterscheiden Messdaten, Studien, Berichte und Meinungen.\n3. Wir zeigen Zusammenhänge statt isolierter Ereignisse.\n4. Wir machen Unsicherheiten sichtbar.\n5. Wir erklären Fachbegriffe verständlich.\n6. Wir suchen den Dialog mit Autorinnen und Autoren.\n7. Wir korrigieren Fehler offen.\n8. Wir kennzeichnen Quellen transparent.\n9. Wir vermeiden unnötige Zuspitzung.\n10. Wir befähigen Menschen, sich selbst ein Urteil zu bilden.""",
    "researchChecklist": """□ Original- oder Primärquelle geprüft\n□ Veröffentlichungsdatum geprüft\n□ Quellentyp bestimmt\n□ Kernaussage verständlich formuliert\n□ Eigentliche Veränderung erkannt\n□ Wirkungskette nachvollziehbar\n□ Betroffene natürliche Systeme benannt\n□ Planetare Grenzen zugeordnet\n□ Gesellschaftliche Relevanz erklärt\n□ Unsicherheiten oder Grenzen benannt\n□ Ansprechpartner und Interviewpotenzial geprüft\n□ Bild und Quellen vollständig""",
    "screenGuide": """• Überschrift klar und sachlich\n• Kurztext in verständlicher Sprache, in der Regel 350–550 Zeichen\n• Fließtext annähernd so gut lesbar wie die Überschrift\n• Keine Logos ohne ausdrückliche Nutzungsfreigabe\n• Stattdessen Quellentyp und Herausgeber transparent nennen\n• Keine öffentliche Prioritätszahl\n• Wirkungskette kurz und nur bei belastbarem Zusammenhang anzeigen\n• Unsicherheiten nicht verschweigen\n• Bild als glaubwürdige Assoziation, nicht als erfundenes Ereignis\n• Ziel: in etwa 30 Sekunden ein neues Verständnis ermöglichen""",
}

EDITORIAL_FIELD_SPECS = (
    ("sourceType", "Quellentyp", 1),
    ("coreChange", "Die eigentliche Veränderung", 3),
    ("questionBehindNews", "Die Frage hinter der Nachricht", 3),
    ("causalChain", "Wirkungskette", 4),
    ("affectedSystems", "Betroffene natürliche Systeme", 3),
    ("planetaryBoundaries", "Betroffene planetare Grenzen", 2),
    ("societalRelevance", "Gesellschaftliche Relevanz", 3),
    ("uncertainties", "Unsicherheiten / Grenzen der Aussage", 3),
    ("interviewPotential", "Interviewpotenzial / mögliche Frage", 3),
    ("screenConnection", "Kurze Zeile für den Infoscreen: Was hängt zusammen?", 2),
)


def read_editorial_file() -> dict[str, Any]:
    if not EDITORIAL_FILE.exists():
        return {"version": 1, "guide": dict(DEFAULT_EDITORIAL_GUIDE), "articles": {}}
    try:
        data = json.loads(EDITORIAL_FILE.read_text(encoding="utf-8"))
        if not isinstance(data, dict):
            raise ValueError
    except (OSError, json.JSONDecodeError, ValueError):
        data = {}
    data.setdefault("version", 1)
    guide = data.setdefault("guide", {})
    for key, value in DEFAULT_EDITORIAL_GUIDE.items():
        guide.setdefault(key, value)
    data.setdefault("articles", {})
    return data


def write_editorial_file(data: dict[str, Any]) -> None:
    EDITORIAL_FILE.parent.mkdir(parents=True, exist_ok=True)
    temp = EDITORIAL_FILE.with_suffix(".json.tmp")
    temp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    temp.replace(EDITORIAL_FILE)


class ScrollableTab(ttk.Frame):
    """Vertikal scrollbarer Inhaltsbereich für lange Studio-Reiter."""

    def __init__(self, parent, padding=0):
        super().__init__(parent)
        self.canvas = tk.Canvas(self, highlightthickness=0, borderwidth=0)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.content = ttk.Frame(self.canvas, padding=padding)
        self.window_id = self.canvas.create_window((0, 0), window=self.content, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")
        self.content.bind("<Configure>", self._sync_scrollregion)
        self.canvas.bind("<Configure>", self._sync_width)
        self.canvas.bind("<Enter>", self._bind_wheel)
        self.canvas.bind("<Leave>", self._unbind_wheel)
        self.canvas.bind("<Prior>", lambda _e: self.canvas.yview_scroll(-1, "pages"))
        self.canvas.bind("<Next>", lambda _e: self.canvas.yview_scroll(1, "pages"))
        self.canvas.bind("<Home>", lambda _e: self.canvas.yview_moveto(0.0))
        self.canvas.bind("<End>", lambda _e: self.canvas.yview_moveto(1.0))
        self.canvas.configure(takefocus=True)

    def _sync_scrollregion(self, _event=None):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def _sync_width(self, event):
        self.canvas.itemconfigure(self.window_id, width=max(event.width, 1))

    def _bind_wheel(self, _event=None):
        self.canvas.focus_set()
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        self.canvas.bind_all("<Button-4>", self._on_linux_wheel_up)
        self.canvas.bind_all("<Button-5>", self._on_linux_wheel_down)

    def _unbind_wheel(self, _event=None):
        self.canvas.unbind_all("<MouseWheel>")
        self.canvas.unbind_all("<Button-4>")
        self.canvas.unbind_all("<Button-5>")

    def _on_mousewheel(self, event):
        delta = -1 if event.delta > 0 else 1
        self.canvas.yview_scroll(delta * 3, "units")

    def _on_linux_wheel_up(self, _event):
        self.canvas.yview_scroll(-3, "units")

    def _on_linux_wheel_down(self, _event):
        self.canvas.yview_scroll(3, "units")


class NewsStudio55(base53.NewsStudio53):
    def __init__(self):
        super().__init__()
        self.title("ZUSTAND News Studio 5.5.0")
        self._replace_widget_text("ZUSTAND News Studio 5.3", "ZUSTAND News Studio 5.5.0")
        ensure_workflow_schema(self.contact_db)
        saved_workflow = read_workflow_file()
        if saved_workflow:
            self.contact_db["workflow"].update(saved_workflow)
        self._save_db()
        write_workflow_file(self.contact_db.get("workflow", {}))
        self.editorial_data = read_editorial_file()
        write_editorial_file(self.editorial_data)
        self._extend_outreach_kinds()
        self._configure_responsive_window()
        self._build_own_finds_tab()
        self._build_workflow_tab()
        self._build_quality_tab()
        self._build_editorial_tab()
        self.status_var.set(
            "News Studio 5.5.0 bereit │ wissenschaftliche Redaktion aktiv"
        )

    def _configure_responsive_window(self) -> None:
        """Passt das Hauptfenster an kleine und große Bildschirme an."""
        self.update_idletasks()
        screen_w = max(self.winfo_screenwidth(), 900)
        screen_h = max(self.winfo_screenheight(), 650)
        width = min(max(int(screen_w * 0.92), 980), 1500)
        height = min(max(int(screen_h * 0.88), 640), 980)
        x = max((screen_w - width) // 2, 0)
        y = max((screen_h - height) // 3, 0)
        self.geometry(f"{width}x{height}+{x}+{y}")
        self.minsize(900, 620)

    def _scrollable_outer(self, tab, padding=14):
        viewport = ScrollableTab(tab, padding=padding)
        viewport.pack(fill="both", expand=True)
        return viewport.content

    # ---------- allgemeine Speicherung ----------
    def _save_db(self) -> None:
        try:
            write_contact_db(self.contact_db)
        except TypeError:
            # Kompatibilität mit älteren Basissignaturen.
            write_contact_db(self.contact_db, getattr(self, "contact_db_path", None))
        # Ältere Studio-Versionen kennen das Workflow-Feld nicht immer.
        write_workflow_file(self.contact_db.get("workflow", {}))

    def _extend_outreach_kinds(self) -> None:
        values = list(self.outreach_kind_combo.cget("values"))
        if "Dankeschön" not in values:
            values.append("Dankeschön")
        self.outreach_kind_combo.configure(values=tuple(values))

    # ---------- Eigene Funde ----------
    def _build_own_finds_tab(self) -> None:
        self.own_finds_tab = ttk.Frame(self.tabs)
        self.tabs.add(self.own_finds_tab, text="Eigene Funde")
        outer = self._scrollable_outer(self.own_finds_tab, padding=14)

        ttk.Label(outer, text="Eigene Fundstücke", font=("Segoe UI", 16, "bold")).pack(anchor="w")
        ttk.Label(
            outer,
            text=(
                "LinkedIn-Link, DOI, Originalquelle, Abstract, PDF oder Screenshot lokal erfassen. "
                "Das Studio erfindet keine Inhalte: Die wissenschaftliche Prüfung und der strukturierte "
                "Beitrag entstehen anschließend aus der Originalquelle."
            ),
            wraplength=1050,
            justify="left",
        ).pack(anchor="w", pady=(2, 10))

        form = ttk.LabelFrame(outer, text="Neuer Fund", padding=10)
        form.pack(fill="x")
        self.find_title_var = tk.StringVar()
        self.find_source_var = tk.StringVar()
        self.find_kind_var = tk.StringVar(value="URL / LinkedIn")
        self.find_file_var = tk.StringVar()
        for label, var in (("Arbeitstitel", self.find_title_var), ("Link / DOI", self.find_source_var)):
            row = ttk.Frame(form); row.pack(fill="x", pady=3)
            ttk.Label(row, text=label, width=18).pack(side="left")
            ttk.Entry(row, textvariable=var).pack(side="left", fill="x", expand=True)
        row = ttk.Frame(form); row.pack(fill="x", pady=3)
        ttk.Label(row, text="Art", width=18).pack(side="left")
        ttk.Combobox(row, textvariable=self.find_kind_var, state="readonly", values=(
            "URL / LinkedIn", "DOI", "PDF", "Screenshot", "Abstract / Text", "Newsletter / E-Mail"
        )).pack(side="left", fill="x", expand=True)
        row = ttk.Frame(form); row.pack(fill="x", pady=3)
        ttk.Label(row, text="Datei", width=18).pack(side="left")
        ttk.Entry(row, textvariable=self.find_file_var).pack(side="left", fill="x", expand=True)
        ttk.Button(row, text="Datei wählen", command=self.choose_find_file).pack(side="left", padx=(6, 0))
        ttk.Label(form, text="Notiz / kopierter Text").pack(anchor="w", pady=(7, 2))
        self.find_text = tk.Text(form, height=7, wrap="word")
        self.find_text.pack(fill="x")
        actions = ttk.Frame(form); actions.pack(fill="x", pady=(8, 0))
        ttk.Button(actions, text="Fund speichern", command=self.save_own_find).pack(side="left")
        ttk.Button(actions, text="Felder leeren", command=self.clear_own_find_form).pack(side="left", padx=6)

        list_frame = ttk.LabelFrame(outer, text="Gespeicherte Funde", padding=8)
        list_frame.pack(fill="both", expand=True, pady=(10, 0))
        cols = ("created", "kind", "title", "status")
        self.find_tree = ttk.Treeview(list_frame, columns=cols, show="headings", height=8)
        for col, title, width in (("created", "Erfasst", 145), ("kind", "Art", 130), ("title", "Titel", 520), ("status", "Status", 160)):
            self.find_tree.heading(col, text=title); self.find_tree.column(col, width=width, anchor="w")
        self.find_tree.pack(fill="both", expand=True)
        bottom = ttk.Frame(list_frame); bottom.pack(fill="x", pady=(6, 0))
        ttk.Button(bottom, text="Recherchepaket kopieren", command=self.copy_find_package).pack(side="left")
        ttk.Button(bottom, text="Als geprüft markieren", command=lambda: self.set_find_status("Geprüft")).pack(side="left", padx=5)
        ttk.Button(bottom, text="Löschen", command=self.delete_own_find).pack(side="right")
        self.refresh_own_finds()

    def choose_find_file(self) -> None:
        if filedialog is None:
            return
        path = filedialog.askopenfilename(title="PDF oder Screenshot auswählen")
        if path:
            self.find_file_var.set(path)

    def clear_own_find_form(self) -> None:
        self.find_title_var.set(""); self.find_source_var.set(""); self.find_file_var.set("")
        self.find_text.delete("1.0", "end")

    def save_own_find(self) -> None:
        title = clean_text(self.find_title_var.get())
        source = clean_text(self.find_source_var.get())
        text = self.find_text.get("1.0", "end").strip()
        file_path = clean_text(self.find_file_var.get())
        if not any((title, source, text, file_path)):
            messagebox.showwarning("Keine Angaben", "Bitte Link, Datei, Titel oder Text angeben.", parent=self)
            return
        find_id = f"FUND_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{safe_slug(title or source)[:20]}"
        stored_file = ""
        if file_path and Path(file_path).exists():
            target_dir = SCRIPT_DIR / "newsredaktion" / "eigene_funde"
            target_dir.mkdir(parents=True, exist_ok=True)
            src = Path(file_path)
            target = target_dir / f"{find_id}_{safe_slug(src.stem)}{src.suffix.lower()}"
            shutil.copy2(src, target)
            stored_file = str(target.relative_to(SCRIPT_DIR))
        self.contact_db["ownFinds"][find_id] = {
            "id": find_id, "createdAt": now_iso(), "updatedAt": now_iso(),
            "kind": self.find_kind_var.get(), "title": title, "source": source,
            "text": text, "file": stored_file, "status": "Neu"
        }
        self._save_db(); self.refresh_own_finds(); self.clear_own_find_form()
        self.status_var.set("Eigenes Fundstück lokal gespeichert")

    def refresh_own_finds(self) -> None:
        if not hasattr(self, "find_tree"): return
        for item in self.find_tree.get_children(): self.find_tree.delete(item)
        finds = self.contact_db.get("ownFinds", {})
        for find_id, item in sorted(finds.items(), key=lambda pair: pair[1].get("createdAt", ""), reverse=True):
            self.find_tree.insert("", "end", iid=find_id, values=(item.get("createdAt", "")[:16].replace("T", " "), item.get("kind", ""), item.get("title") or item.get("source") or "Ohne Titel", item.get("status", "Neu")))

    def _selected_find(self) -> tuple[str, dict[str, Any]] | tuple[None, None]:
        selection = self.find_tree.selection()
        if not selection: return None, None
        find_id = selection[0]
        return find_id, self.contact_db.get("ownFinds", {}).get(find_id)

    def copy_find_package(self) -> None:
        _id, item = self._selected_find()
        if not item:
            messagebox.showinfo("Auswahl fehlt", "Bitte zuerst einen Fund auswählen.", parent=self); return
        package = f"""Bitte prüfe dieses Fundstück für den ZUSTAND-Infoscreen und erstelle daraus einen strukturierten Studio-Datensatz.\n\nArt: {item.get('kind','')}\nArbeitstitel: {item.get('title','')}\nLink/DOI: {item.get('source','')}\nLokale Datei: {item.get('file','')}\n\nNotiz/Text:\n{item.get('text','')}\n\nBitte zuerst die Original- bzw. Primärquelle verifizieren. Danach: Kurztext (350–550 Zeichen), Quellentyp, eigentliche Veränderung, Wirkungskette, betroffene natürliche Systeme und planetare Grenzen, gesellschaftliche Relevanz, Quellen, Bildidee, Unsicherheiten, Interviewpotenzial und geeignete Ansprechperson. Keine öffentliche Prioritätszahl verwenden."""
        self.clipboard_clear(); self.clipboard_append(package); self.update()
        self.status_var.set("Recherchepaket in die Zwischenablage kopiert")

    def set_find_status(self, status: str) -> None:
        find_id, item = self._selected_find()
        if not item: return
        item["status"] = status; item["updatedAt"] = now_iso(); self._save_db(); self.refresh_own_finds()

    def delete_own_find(self) -> None:
        find_id, item = self._selected_find()
        if not item: return
        if not messagebox.askyesno("Fund löschen", "Dieses lokale Fundstück wirklich löschen?", parent=self): return
        file_ref = clean_text(item.get("file"))
        if file_ref:
            path = SCRIPT_DIR / file_ref
            try:
                if path.exists(): path.unlink()
            except OSError: pass
        self.contact_db["ownFinds"].pop(find_id, None); self._save_db(); self.refresh_own_finds()

    # ---------- Workflow, Status, Statistik, Löschen ----------
    def _build_workflow_tab(self) -> None:
        self.workflow_tab = ttk.Frame(self.tabs)
        self.tabs.add(self.workflow_tab, text="Workflow")
        outer = self._scrollable_outer(self.workflow_tab, padding=14)
        ttk.Label(outer, text="Kontakt- und Interviewworkflow", font=("Segoe UI", 16, "bold")).pack(anchor="w")

        stats = ttk.LabelFrame(outer, text="Redaktionsstatistik", padding=8); stats.pack(fill="x", pady=(8, 10))
        self.stats_var = tk.StringVar(); ttk.Label(stats, textvariable=self.stats_var, font=("Segoe UI", 11, "bold")).pack(anchor="w")

        body = ttk.Frame(outer); body.pack(fill="both", expand=True)
        body.columnconfigure(0, weight=3)
        body.columnconfigure(1, weight=2)
        body.rowconfigure(0, weight=1)
        left = ttk.LabelFrame(body, text="Kontakte", padding=8); left.grid(row=0, column=0, sticky="nsew")
        tree_wrap = ttk.Frame(left); tree_wrap.pack(fill="both", expand=True)
        workflow_scroll = ttk.Scrollbar(tree_wrap, orient="vertical")
        self.workflow_tree = ttk.Treeview(tree_wrap, columns=("name", "institution", "status", "links"), show="headings", height=12, yscrollcommand=workflow_scroll.set)
        workflow_scroll.configure(command=self.workflow_tree.yview)
        for col, title, width in (("name", "Name", 210), ("institution", "Institution", 270), ("status", "Status", 180), ("links", "Beiträge", 70)):
            self.workflow_tree.heading(col, text=title); self.workflow_tree.column(col, width=width, anchor="w")
        self.workflow_tree.pack(side="left", fill="both", expand=True)
        workflow_scroll.pack(side="right", fill="y")
        self.workflow_tree.bind("<<TreeviewSelect>>", self._load_workflow_contact)

        right = ttk.LabelFrame(body, text="Bearbeitung", padding=10); right.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        self.workflow_name_var = tk.StringVar(value="Kein Kontakt gewählt")
        ttk.Label(right, textvariable=self.workflow_name_var, font=("Segoe UI", 11, "bold"), wraplength=330).pack(anchor="w")
        ttk.Label(right, text="Status").pack(anchor="w", pady=(10, 2))
        self.workflow_status_var = tk.StringVar(value=STATUS_VALUES[0])
        ttk.Combobox(right, textvariable=self.workflow_status_var, values=STATUS_VALUES, state="readonly", width=34).pack(anchor="w")
        ttk.Button(right, text="Status speichern", command=self.save_contact_status).pack(anchor="w", pady=6)
        ttk.Label(right, text="Notizen").pack(anchor="w", pady=(8, 2))
        self.workflow_notes = tk.Text(right, width=42, height=8, wrap="word"); self.workflow_notes.pack(fill="x")
        ttk.Button(right, text="Notizen speichern", command=self.save_workflow_notes).pack(anchor="w", pady=6)
        ttk.Separator(right).pack(fill="x", pady=8)
        ttk.Button(right, text="Interviewmappe erstellen", command=self.create_interview_map).pack(fill="x", pady=2)
        ttk.Button(right, text="Artikel-Verknüpfung entfernen", command=self.remove_article_link_dialog).pack(fill="x", pady=2)
        ttk.Button(right, text="Verwaiste Kontakte prüfen", command=self.cleanup_orphan_contacts).pack(fill="x", pady=2)
        self.refresh_workflow()

    def refresh_workflow(self, keep_selected: str | None = None) -> None:
        ensure_workflow_schema(self.contact_db)
        if hasattr(self, "workflow_tree"):
            if keep_selected is None:
                selection = self.workflow_tree.selection()
                keep_selected = selection[0] if selection else None
            for item in self.workflow_tree.get_children():
                self.workflow_tree.delete(item)
            for cid, contact in sorted(
                self.contact_db.get("contacts", {}).items(),
                key=lambda p: clean_text(p[1].get("name")).lower(),
            ):
                links = sum(
                    1
                    for link in self.contact_db.get("articleLinks", {}).values()
                    if cid in link.get("contactIds", [])
                )
                wf = self.contact_db["workflow"].get(cid, {})
                self.workflow_tree.insert(
                    "",
                    "end",
                    iid=cid,
                    values=(
                        clean_text(contact.get("name")) or "Ohne Namen",
                        clean_text(contact.get("institution")),
                        wf.get("status", "Recherchiert"),
                        links,
                    ),
                )
            if keep_selected and self.workflow_tree.exists(keep_selected):
                self.workflow_tree.selection_set(keep_selected)
                self.workflow_tree.focus(keep_selected)
                self.workflow_tree.see(keep_selected)
        contacts = len(self.contact_db.get("contacts", {})); articles = len(self.contact_db.get("articleLinks", {})); finds = len(self.contact_db.get("ownFinds", {}))
        counts = {status: 0 for status in STATUS_VALUES}
        for wf in self.contact_db.get("workflow", {}).values():
            status = wf.get("status", "Recherchiert"); counts[status] = counts.get(status, 0) + 1
        if hasattr(self, "stats_var"):
            self.stats_var.set(f"Beiträge {articles}   │   Kontakte {contacts}   │   Eigene Funde {finds}   │   Mails {counts.get('E-Mail versendet',0)}   │   Antworten {counts.get('Antwort erhalten',0)}   │   Interviews {counts.get('Interview durchgeführt',0)}")
        if hasattr(self, "refresh_outreach_contacts"): self.refresh_outreach_contacts()

    def _selected_workflow_contact(self) -> tuple[str, dict[str, Any]] | tuple[None, None]:
        selection = self.workflow_tree.selection()
        if not selection: return None, None
        cid = selection[0]; return cid, self.contact_db.get("contacts", {}).get(cid)

    def _load_workflow_contact(self, _event=None) -> None:
        cid, contact = self._selected_workflow_contact()
        if not contact: return
        wf = self.contact_db["workflow"].setdefault(cid, {"status": "Recherchiert", "history": [], "notes": ""})
        self.workflow_name_var.set(clean_text(contact.get("name")) or clean_text(contact.get("institution")) or cid)
        self.workflow_status_var.set(wf.get("status", "Recherchiert"))
        self.workflow_notes.delete("1.0", "end"); self.workflow_notes.insert("1.0", wf.get("notes", ""))

    def save_contact_status(self) -> None:
        cid, _contact = self._selected_workflow_contact()
        if not cid: return
        wf = self.contact_db["workflow"].setdefault(cid, {"history": [], "notes": ""})
        old = wf.get("status", "Recherchiert"); new = self.workflow_status_var.get()
        wf["status"] = new
        wf["updatedAt"] = now_iso()
        if new != old:
            wf.setdefault("history", []).append(
                {"at": now_iso(), "from": old, "to": new}
            )
        self._save_db()
        self.refresh_workflow(keep_selected=cid)
        self._load_workflow_contact()
        self.status_var.set(f"Kontaktstatus dauerhaft gespeichert: {new}")

    def save_workflow_notes(self) -> None:
        cid, _contact = self._selected_workflow_contact()
        if not cid: return
        wf = self.contact_db["workflow"].setdefault(cid, {"status": "Recherchiert", "history": []})
        wf["notes"] = self.workflow_notes.get("1.0", "end").strip(); wf["updatedAt"] = now_iso()
        self._save_db(); self.status_var.set("Workflow-Notizen gespeichert")

    def remove_article_link_dialog(self) -> None:
        cid, contact = self._selected_workflow_contact()
        if not cid or not contact:
            messagebox.showinfo("Auswahl fehlt", "Bitte einen Kontakt auswählen.", parent=self); return
        linked = [(key, link) for key, link in self.contact_db.get("articleLinks", {}).items() if cid in link.get("contactIds", [])]
        if not linked:
            messagebox.showinfo("Keine Verknüpfung", "Dieser Kontakt ist mit keinem Beitrag verknüpft.", parent=self); return
        dialog = tk.Toplevel(self); dialog.title("Artikel-Verknüpfung entfernen"); dialog.transient(self); dialog.grab_set()
        ttk.Label(dialog, text="Nur die gewählte Verknüpfung wird entfernt. Der Kontakt bleibt zunächst erhalten.", wraplength=600).pack(anchor="w", padx=12, pady=10)
        var = tk.StringVar(value=linked[0][0]); combo = ttk.Combobox(dialog, textvariable=var, state="readonly", width=80, values=[key for key, _ in linked]); combo.pack(padx=12, pady=5)
        label_map = {key: clean_text(link.get("title")) or key for key, link in linked}
        info = tk.StringVar(value=label_map[var.get()]); ttk.Label(dialog, textvariable=info, wraplength=650).pack(padx=12, pady=4)
        combo.bind("<<ComboboxSelected>>", lambda _e: info.set(label_map.get(var.get(), var.get())))
        def do_remove():
            key = var.get(); link = self.contact_db.get("articleLinks", {}).get(key, {})
            link["contactIds"] = [x for x in link.get("contactIds", []) if x != cid]
            if not link["contactIds"]: self.contact_db["articleLinks"].pop(key, None)
            self._save_db(); dialog.destroy(); self.refresh_workflow()
            self.status_var.set("Artikel-Kontakt-Verknüpfung entfernt; Kontakt wurde nicht automatisch gelöscht")
        ttk.Button(dialog, text="Verknüpfung entfernen", command=do_remove).pack(pady=12)

    def cleanup_orphan_contacts(self) -> None:
        used = set()
        for link in self.contact_db.get("articleLinks", {}).values(): used.update(link.get("contactIds", []))
        orphans = [cid for cid in self.contact_db.get("contacts", {}) if cid not in used]
        if not orphans:
            messagebox.showinfo("Prüfung abgeschlossen", "Es gibt keine verwaisten Kontakte.", parent=self); return
        names = [clean_text(self.contact_db["contacts"][cid].get("name")) or cid for cid in orphans]
        answer = messagebox.askyesno("Verwaiste Kontakte", f"{len(orphans)} Kontakt(e) haben keine Beitragsverknüpfung:\n\n" + "\n".join(names[:12]) + ("\n…" if len(names)>12 else "") + "\n\nDiese Kontakte jetzt löschen?", parent=self)
        if not answer: return
        for cid in orphans:
            self.contact_db["contacts"].pop(cid, None); self.contact_db.get("workflow", {}).pop(cid, None)
        self._save_db(); self.refresh_workflow(); self.status_var.set(f"{len(orphans)} verwaiste Kontakte gelöscht")

    # ---------- Qualitätsprüfung ----------
    def _build_quality_tab(self) -> None:
        self.quality_tab = ttk.Frame(self.tabs); self.tabs.add(self.quality_tab, text="Qualitätsprüfung")
        outer = self._scrollable_outer(self.quality_tab, padding=14)
        ttk.Label(outer, text="Redaktionsprüfung vor Veröffentlichung", font=("Segoe UI", 16, "bold")).pack(anchor="w")
        row = ttk.Frame(outer); row.pack(fill="x", pady=(10, 8))
        ttk.Label(row, text="Beitrag", width=14).pack(side="left")
        self.quality_article_var = tk.StringVar(); self.quality_article_combo = ttk.Combobox(row, textvariable=self.quality_article_var, state="readonly")
        self.quality_article_combo.pack(side="left", fill="x", expand=True); self.quality_article_combo.bind("<<ComboboxSelected>>", self.load_quality_check)
        ttk.Button(row, text="Aktualisieren", command=self.refresh_quality_articles).pack(side="left", padx=5)
        self.quality_vars: dict[str, Any] = {}
        checks = ttk.LabelFrame(outer, text="Checkliste", padding=10); checks.pack(fill="x")
        for key, label in QUALITY_FIELDS:
            var = tk.BooleanVar(value=False); self.quality_vars[key] = var
            ttk.Checkbutton(checks, text=label, variable=var).pack(anchor="w", pady=3)
        ttk.Label(outer, text="Redaktionelle Notiz").pack(anchor="w", pady=(10, 2))
        self.quality_notes = tk.Text(outer, height=6, wrap="word"); self.quality_notes.pack(fill="x")
        actions = ttk.Frame(outer); actions.pack(fill="x", pady=8)
        ttk.Button(actions, text="Prüfung speichern", command=self.save_quality_check).pack(side="left")
        ttk.Button(actions, text="Freigabe prüfen", command=self.check_release_readiness).pack(side="left", padx=6)
        self.quality_ids: list[str] = []; self.refresh_quality_articles()

    def refresh_quality_articles(self) -> None:
        links = self.contact_db.get("articleLinks", {})
        ordered = sorted(links.items(), key=lambda p: clean_text(p[1].get("title")).lower())
        self.quality_ids = [key for key, _ in ordered]
        self.quality_article_combo["values"] = [clean_text(link.get("title")) or key for key, link in ordered]
        if self.quality_ids:
            self.quality_article_combo.current(0); self.load_quality_check()

    def load_quality_check(self, _event=None) -> None:
        idx = self.quality_article_combo.current()
        if idx < 0 or idx >= len(self.quality_ids): return
        check = self.contact_db.get("qualityChecks", {}).get(self.quality_ids[idx], {})
        for key, _label in QUALITY_FIELDS: self.quality_vars[key].set(bool(check.get(key, False)))
        self.quality_notes.delete("1.0", "end"); self.quality_notes.insert("1.0", check.get("notes", ""))

    def save_quality_check(self) -> None:
        idx = self.quality_article_combo.current()
        if idx < 0 or idx >= len(self.quality_ids): return
        key = self.quality_ids[idx]
        data = {field: bool(var.get()) for field, var in self.quality_vars.items()}
        data.update({"notes": self.quality_notes.get("1.0", "end").strip(), "updatedAt": now_iso()})
        self.contact_db["qualityChecks"][key] = data; self._save_db(); self.status_var.set("Qualitätsprüfung gespeichert")

    def check_release_readiness(self) -> None:
        missing = [label for key, label in QUALITY_FIELDS if not self.quality_vars[key].get()]
        if missing:
            messagebox.showwarning("Noch nicht freigabereif", "Folgende Punkte fehlen:\n\n• " + "\n• ".join(missing), parent=self)
        else:
            messagebox.showinfo("Freigabereif", "Alle Qualitätskriterien sind bestätigt.", parent=self)

    # ---------- Wissenschaftliche Redaktion ----------
    def _build_editorial_tab(self) -> None:
        self.editorial_tab = ttk.Frame(self.tabs)
        self.tabs.add(self.editorial_tab, text="Redaktion")
        outer = self._scrollable_outer(self.editorial_tab, padding=14)

        ttk.Label(
            outer,
            text="Wissenschaftliche Redaktion",
            font=("Segoe UI", 16, "bold"),
        ).pack(anchor="w")
        ttk.Label(
            outer,
            text=(
                "Hier werden Leitbild, Redaktionsregeln und die Zusammenhänge hinter "
                "den Meldungen dauerhaft festgehalten. Die Angaben bleiben lokal und "
                "werden noch nicht automatisch auf dem Infoscreen veröffentlicht."
            ),
            wraplength=1100,
            justify="left",
        ).pack(anchor="w", pady=(2, 10))

        guide_box = ttk.LabelFrame(outer, text="Leitbild und Redaktionshandbuch", padding=10)
        guide_box.pack(fill="x")
        self.editorial_guide_widgets: dict[str, tk.Text] = {}
        guide_specs = (
            ("mission", "Leitbild", 4),
            ("code", "Redaktionskodex", 11),
            ("researchChecklist", "Recherche-Checkliste", 12),
            ("screenGuide", "Gestaltung des Infoscreens", 11),
        )
        for key, label, height in guide_specs:
            ttk.Label(guide_box, text=label, font=("Segoe UI", 10, "bold")).pack(
                anchor="w", pady=(8 if self.editorial_guide_widgets else 0, 2)
            )
            widget = tk.Text(guide_box, height=height, wrap="word", undo=True)
            widget.pack(fill="x")
            widget.insert("1.0", self.editorial_data["guide"].get(key, ""))
            self.editorial_guide_widgets[key] = widget
        guide_actions = ttk.Frame(guide_box)
        guide_actions.pack(fill="x", pady=(8, 0))
        ttk.Button(
            guide_actions,
            text="Redaktionshandbuch speichern",
            command=self.save_editorial_guide,
        ).pack(side="left")
        ttk.Button(
            guide_actions,
            text="Standardtexte wiederherstellen",
            command=self.restore_editorial_guide,
        ).pack(side="left", padx=6)

        article_box = ttk.LabelFrame(outer, text="Zusammenhänge zu einem Beitrag", padding=10)
        article_box.pack(fill="x", pady=(12, 0))
        select_row = ttk.Frame(article_box)
        select_row.pack(fill="x", pady=(0, 8))
        ttk.Label(select_row, text="Beitrag", width=18).pack(side="left")
        self.editorial_article_var = tk.StringVar()
        self.editorial_article_combo = ttk.Combobox(
            select_row, textvariable=self.editorial_article_var, state="readonly"
        )
        self.editorial_article_combo.pack(side="left", fill="x", expand=True)
        self.editorial_article_combo.bind(
            "<<ComboboxSelected>>", self.load_editorial_article
        )
        ttk.Button(
            select_row, text="Liste aktualisieren", command=self.refresh_editorial_articles
        ).pack(side="left", padx=(6, 0))

        self.editorial_field_widgets: dict[str, tk.Text] = {}
        for key, label, height in EDITORIAL_FIELD_SPECS:
            ttk.Label(article_box, text=label, font=("Segoe UI", 10, "bold")).pack(
                anchor="w", pady=(7, 2)
            )
            widget = tk.Text(article_box, height=height, wrap="word", undo=True)
            widget.pack(fill="x")
            self.editorial_field_widgets[key] = widget

        action_row = ttk.Frame(article_box)
        action_row.pack(fill="x", pady=(9, 0))
        ttk.Button(
            action_row,
            text="Redaktionelle Angaben speichern",
            command=self.save_editorial_article,
        ).pack(side="left")
        ttk.Button(
            action_row,
            text="Angaben leeren",
            command=self.clear_editorial_article_fields,
        ).pack(side="left", padx=6)
        ttk.Button(
            action_row,
            text="Rechercheauftrag kopieren",
            command=self.copy_editorial_research_prompt,
        ).pack(side="right")

        self.editorial_article_ids: list[str] = []
        self.refresh_editorial_articles()

    def save_editorial_guide(self) -> None:
        for key, widget in self.editorial_guide_widgets.items():
            self.editorial_data["guide"][key] = widget.get("1.0", "end").strip()
        self.editorial_data["updatedAt"] = now_iso()
        write_editorial_file(self.editorial_data)
        self.status_var.set("Redaktionshandbuch dauerhaft gespeichert")

    def restore_editorial_guide(self) -> None:
        if not messagebox.askyesno(
            "Standardtexte wiederherstellen",
            "Die bearbeiteten Texte des Redaktionshandbuchs ersetzen?",
            parent=self,
        ):
            return
        self.editorial_data["guide"] = dict(DEFAULT_EDITORIAL_GUIDE)
        for key, widget in self.editorial_guide_widgets.items():
            widget.delete("1.0", "end")
            widget.insert("1.0", self.editorial_data["guide"].get(key, ""))
        write_editorial_file(self.editorial_data)
        self.status_var.set("Standardtexte des Redaktionshandbuchs wiederhergestellt")

    def refresh_editorial_articles(self) -> None:
        links = self.contact_db.get("articleLinks", {})
        ordered = sorted(
            ((key, value) for key, value in links.items() if isinstance(value, dict)),
            key=lambda pair: clean_text(pair[1].get("title")).lower(),
        )
        previous_id = self._selected_editorial_article_id()
        self.editorial_article_ids = [key for key, _value in ordered]
        labels = [clean_text(value.get("title")) or key for key, value in ordered]
        self.editorial_article_combo["values"] = labels
        if not labels:
            self.editorial_article_var.set("")
            self.clear_editorial_article_fields()
            return
        if previous_id in self.editorial_article_ids:
            index = self.editorial_article_ids.index(previous_id)
        else:
            index = 0
        self.editorial_article_combo.current(index)
        self.load_editorial_article()

    def _selected_editorial_article_id(self) -> str | None:
        if not hasattr(self, "editorial_article_combo"):
            return None
        index = self.editorial_article_combo.current()
        if 0 <= index < len(getattr(self, "editorial_article_ids", [])):
            return self.editorial_article_ids[index]
        return None

    def clear_editorial_article_fields(self) -> None:
        for widget in getattr(self, "editorial_field_widgets", {}).values():
            widget.delete("1.0", "end")

    def load_editorial_article(self, _event=None) -> None:
        article_id = self._selected_editorial_article_id()
        self.clear_editorial_article_fields()
        if not article_id:
            return
        data = self.editorial_data.get("articles", {}).get(article_id, {})
        for key, _label, _height in EDITORIAL_FIELD_SPECS:
            self.editorial_field_widgets[key].insert("1.0", clean_text(data.get(key)))

    def save_editorial_article(self) -> None:
        article_id = self._selected_editorial_article_id()
        if not article_id:
            messagebox.showinfo(
                "Beitrag fehlt", "Bitte zuerst einen Beitrag auswählen.", parent=self
            )
            return
        link = self.contact_db.get("articleLinks", {}).get(article_id, {})
        values = {
            key: widget.get("1.0", "end").strip()
            for key, widget in self.editorial_field_widgets.items()
        }
        values.update(
            {
                "articleId": article_id,
                "title": clean_text(link.get("title")),
                "sourceUrl": clean_text(link.get("sourceUrl")),
                "updatedAt": now_iso(),
            }
        )
        self.editorial_data.setdefault("articles", {})[article_id] = values
        self.editorial_data["updatedAt"] = now_iso()
        write_editorial_file(self.editorial_data)
        self.status_var.set("Redaktionelle Angaben zum Beitrag dauerhaft gespeichert")

    def copy_editorial_research_prompt(self) -> None:
        article_id = self._selected_editorial_article_id()
        if not article_id:
            messagebox.showinfo(
                "Beitrag fehlt", "Bitte zuerst einen Beitrag auswählen.", parent=self
            )
            return
        link = self.contact_db.get("articleLinks", {}).get(article_id, {})
        existing = self.editorial_data.get("articles", {}).get(article_id, {})
        prompt = f"""Prüfe die folgende Meldung für den ZUSTAND-Infoscreen ausschließlich anhand seriöser Original- oder Primärquellen.

Titel: {clean_text(link.get('title'))}
Originalquelle: {clean_text(link.get('sourceUrl'))}

Arbeite besonders heraus:
1. Welche tatsächliche Veränderung steckt hinter der Nachricht?
2. Was erklärt sie über das Funktionieren unserer natürlichen Lebensgrundlagen?
3. Welche belastbare Ursache-Wirkungs-Kette lässt sich kurz darstellen?
4. Welche natürlichen Systeme und planetaren Grenzen sind betroffen?
5. Warum ist das gesellschaftlich und gesundheitlich relevant?
6. Welche Unsicherheiten, Grenzen oder alternativen Erklärungen müssen genannt werden?
7. Welcher Quellentyp liegt vor: Messdaten, Studie, Bericht, Behörde, Universität, UN-Organisation oder anderes?
8. Welche kurze Zeile „Was hängt zusammen?“ eignet sich für den Infoscreen?
9. Gibt es geeignete Autorinnen, Autoren oder Forschende für ein Interview?

Vorhandene redaktionelle Notizen:
{json.dumps(existing, ensure_ascii=False, indent=2)}

Formuliere sachlich, leicht verständlich und ohne Alarmismus. Keine öffentliche Prioritätszahl verwenden."""
        self.clipboard_clear()
        self.clipboard_append(prompt)
        self.update()
        self.status_var.set("Redaktioneller Rechercheauftrag kopiert")

    # ---------- Interviewtexte und Interviewmappe ----------
    def generate_outreach_text(self) -> None:
        _cid, contact = self._selected_contact()
        if not contact: return
        article = self._selected_article(); title = article.get("title") or "[Titel der Veröffentlichung]"
        greeting = base53.greeting_for(contact); hook = clean_text(self.outreach_hook_var.get()); hook_paragraph = f"\n\n{hook}" if hook else ""
        kind = self.outreach_kind_var.get()
        studio = "Die Studios des Offenen Kanals Lübeck sind montags bis freitags von 12:00 bis 19:00 Uhr geöffnet. Innerhalb dieses Zeitrahmens richten wir uns gern nach Ihren Möglichkeiten."
        sequence = "Geplant sind eine kurze Vorstellung Ihrer Person und Arbeit, die verständliche Einordnung der Veröffentlichung und anschließend meine Fragen. Bei einer späteren Live-Sendung könnte das Telefon zusätzlich für Fragen von Hörerinnen und Hörern geöffnet werden; dies würden wir vorher ausdrücklich mit Ihnen abstimmen."
        precheck = "Vor dem Gespräch erhalten Sie den vorbereiteten Infoscreen-Beitrag und die vorgesehenen Fragen zur Durchsicht. Hinweise oder Korrekturen berücksichtigen wir selbstverständlich vor der Aufzeichnung."

        if kind == "Dankeschön":
            subject = f"Vielen Dank für das Gespräch zu „{title}“"
            body = f"""{greeting}\n\nherzlichen Dank, dass Sie sich die Zeit für unser Gespräch über Ihre Veröffentlichung „{title}“ genommen haben. Ihre Erläuterungen helfen sehr dabei, die Forschung verständlich und zugleich differenziert zu vermitteln.\n\nSobald der Beitrag beziehungsweise die Aufzeichnung veröffentlicht ist, sende ich Ihnen den Link. Über Hinweise oder Korrekturen freue ich mich auch im Nachgang.\n\nMit freundlichen Grüßen\n\nDetlef Hau\nTechnische Hochschule Lübeck\nProjekt ZUSTAND – Die Vermessung unserer Zukunft"""
        elif kind == "Freundliche Erinnerung":
            subject = "Freundliche Nachfrage zu meiner Interviewanfrage"
            body = f"""{greeting}\n\nvor einigen Tagen hatte ich Sie wegen eines kurzen Gesprächs zu Ihrer Veröffentlichung „{title}“ angeschrieben.\n\nDa Ihre Arbeit sehr gut zu unserem Bildungs- und Öffentlichkeitsprojekt ZUSTAND – Die Vermessung unserer Zukunft passt, möchte ich freundlich nachfragen, ob grundsätzlich Interesse an einem etwa 20- bis 30-minütigen Telefon- oder Online-Interview besteht.{hook_paragraph}\n\n{studio}\n\nFalls es derzeit zeitlich nicht passt, genügt selbstverständlich eine kurze Rückmeldung.\n\nMit freundlichen Grüßen\n\nDetlef Hau\nTechnische Hochschule Lübeck\nProjekt ZUSTAND – Die Vermessung unserer Zukunft"""
        elif kind == "Intervieweinladung":
            subject = f"Abstimmung unseres Gesprächs zu „{title}“"
            body = f"""{greeting}\n\nvielen Dank für Ihre Bereitschaft zu einem Gespräch über Ihre Veröffentlichung „{title}“.\n\nVorgesehen ist ein etwa 20- bis 30-minütiges Telefon- oder Online-Interview.{hook_paragraph}\n\n{sequence}\n\n{studio}\n\n{precheck}\n\nWelche Termine würden Ihnen gut passen?\n\nMit freundlichen Grüßen\n\nDetlef Hau\nTechnische Hochschule Lübeck\nProjekt ZUSTAND – Die Vermessung unserer Zukunft"""
        elif kind == "Telefonleitfaden":
            subject = f"Telefonleitfaden: {title}"
            body = f"""KONTAKT\n{clean_text(contact.get('name'))}\n{clean_text(contact.get('role'))}\n{clean_text(contact.get('institution'))}\n{clean_text(contact.get('phone'))}\n\nBEGRÜSSUNG\nGuten Tag, mein Name ist Detlef Hau von der Technischen Hochschule Lübeck. Passt es gerade für eine kurze Frage?\n\nANLASS\nIch beschäftige mich im Bildungs- und Öffentlichkeitsprojekt ZUSTAND – Die Vermessung unserer Zukunft mit aktuellen wissenschaftlichen Erkenntnissen zu Umwelt, Gesellschaft und Zukunftsfähigkeit.\n\nWARUM DIESE PERSON?\nBei unserer Recherche bin ich auf die Veröffentlichung „{title}“ gestoßen.\n{hook or 'Die Ergebnisse erscheinen besonders geeignet, um sie einer breiteren Öffentlichkeit verständlich zu vermitteln.'}\n\nANFRAGE\nHätten Sie grundsätzlich Interesse an einem etwa 20- bis 30-minütigen Telefon- oder Online-Interview?\n\nABLAUF\n{sequence}\n\nZEITRAHMEN\n{studio}\n\nTRANSPARENZ\n{precheck}\n\nNOTIZEN\n• Interesse:\n• Bevorzugtes Format:\n• Mögliche Termine:\n• Hörerfragen möglich?\n• Bedingungen/Wünsche:\n• Nächster Schritt:"""
        else:
            subject = "Anfrage zu einem kurzen Gespräch über Ihre aktuelle Veröffentlichung"
            body = f"""{greeting}\n\nich bin Lehrbeauftragter an der Technischen Hochschule Lübeck und beschäftige mich im Bildungs- und Öffentlichkeitsprojekt ZUSTAND – Die Vermessung unserer Zukunft mit aktuellen wissenschaftlichen Erkenntnissen zu Umwelt, Gesellschaft und Zukunftsfähigkeit.\n\nBei der Recherche für unseren öffentlichen Infoscreen bin ich auf Ihre Veröffentlichung „{title}“ aufmerksam geworden. Die Ergebnisse erscheinen mir besonders relevant und verständlich für eine breite Öffentlichkeit.{hook_paragraph}\n\nDeshalb möchte ich Sie fragen, ob Sie Interesse an einem etwa 20- bis 30-minütigen Telefon- oder Online-Interview hätten. Ziel ist es, Ihre Forschung allgemeinverständlich vorzustellen und die wichtigsten Erkenntnisse direkt von Ihnen erläutern zu lassen. Das Gespräch soll unter anderem über den Offenen Kanal Lübeck sowie in unserer Bildungsarbeit verwendet werden.\n\n{sequence}\n\n{studio}\n\n{precheck}\n\nIch würde mich sehr freuen, wenn Sie Zeit für ein Gespräch finden.\n\nMit freundlichen Grüßen\n\nDetlef Hau\nTechnische Hochschule Lübeck\nProjekt ZUSTAND – Die Vermessung unserer Zukunft"""
        self.outreach_subject_var.set(subject); self._set_outreach_text(body)

    def create_interview_map(self) -> None:
        cid, contact = self._selected_workflow_contact()
        if not contact:
            messagebox.showinfo("Auswahl fehlt", "Bitte einen Kontakt auswählen.", parent=self); return
        try:
            from docx import Document
        except ImportError:
            messagebox.showerror("Zusatzmodul fehlt", "Für DOCX-Interviewmappen bitte 'python-docx' installieren.", parent=self); return
        articles = base53.linked_articles(self.contact_db, cid)
        wf = self.contact_db.get("workflow", {}).get(cid, {})
        doc = Document(); doc.add_heading("Interviewmappe – ZUSTAND", 0)
        doc.add_heading(clean_text(contact.get("name")) or "Gesprächspartner*in", 1)
        for label, key in (("Institution", "institution"), ("Funktion", "role"), ("E-Mail", "email"), ("Telefon", "phone"), ("Profil", "profileUrl"), ("ORCID", "orcid")):
            value = clean_text(contact.get(key))
            if value: doc.add_paragraph(f"{label}: {value}")
        doc.add_heading("Verknüpfte Veröffentlichungen", 1)
        if articles:
            for article in articles: doc.add_paragraph(article.get("title", ""), style="List Bullet")
        else: doc.add_paragraph("Noch keine Veröffentlichung verknüpft.")
        doc.add_heading("Geplanter Gesprächsablauf", 1)
        for item in ("Kurze Vorstellung der Person und ihres Arbeitsgebiets", "Verständliche Vorstellung der Veröffentlichung", "Einordnung der wichtigsten Ergebnisse und Grenzen", "Bedeutung für Gesellschaft, Politik, Bildung und Praxis", "Optional und nur nach Absprache: Hörerfragen", "Hinweise auf weiterführende Informationen"):
            doc.add_paragraph(item, style="List Number")
        doc.add_heading("Interviewfragen", 1)
        questions = (
            "Was war der Ausgangspunkt Ihrer Untersuchung?", "Was ist das wichtigste Ergebnis?",
            "Welcher Befund hat Sie selbst besonders überrascht?", "Wie belastbar sind die Ergebnisse und wo liegen Grenzen?",
            "Was wird in der öffentlichen Debatte häufig missverstanden?", "Welche Folgen ergeben sich für Politik oder Praxis?",
            "Welche Rolle spielen soziale Gerechtigkeit und unterschiedliche Betroffenheit?", "Was müsste als Nächstes erforscht werden?",
            "Welche Entwicklung macht Ihnen Sorge?", "Welche realistischen Handlungsmöglichkeiten sehen Sie?"
        )
        for q in questions: doc.add_paragraph(q, style="List Number")
        doc.add_heading("Eigene Notizen", 1); doc.add_paragraph(wf.get("notes", "")); doc.add_paragraph("\n" * 8)
        target_dir = SCRIPT_DIR / "newsredaktion" / "interviewmappen"; target_dir.mkdir(parents=True, exist_ok=True)
        filename = f"interviewmappe_{safe_slug(clean_text(contact.get('name')), cid)}_{datetime.now().strftime('%Y%m%d')}.docx"
        path = target_dir / filename; doc.save(path)
        self.status_var.set(f"Interviewmappe erstellt: {path}")
        messagebox.showinfo("Interviewmappe erstellt", str(path), parent=self)


if __name__ == "__main__":
    app = NewsStudio55()
    app.mainloop()
